from docx import Document
from docx.shared import Inches
import json

#Abre el archivo con la informacion del documento a COmbinar
with open('data2.json') as file:
    records = json.load(file)

archivoOriginal = records[0]["ARCHIVO_ORIGINAL"]
archivoFinal    = records[0]["ARCHIVO_FINAL"]
radicado        = records[0]["RADICADO"]

document = Document(archivoOriginal)

# print (document.tables[1].rows[0].cells[0].text)
# Realiza verficicacion de tablas a combinar...

i=0

for key in document.tables:
    if key:
        valorColumna = document.tables[i].rows[0].cells[0].text
        if valorColumna=='TERCEROS_TABLA':
            tableCombina=document.tables[i]
            nombreCampos = tableCombina.rows[2].cells
            tableIdCombina = i
        i=i+1    
        


i=0
arrNombre = []
for nombre in nombreCampos:
    #arrNombre[i] = 
    arrNombre.append(nombre.text)
   
i=2
# Recorre el array con los campos a combinar.
for datos in records:
    # si i = 2 significa que es el primer elemento de datos y este debe tener los datos del archivo.
    if i>2:
        iC=0
    #   Verifica si existe la celda i qeu inicia en 2 sin no existe la crea.    
        try:
            tableCombina.rows[i-1].cells
        except:
            nuevaRowCell = tableCombina.add_row().cells
            for nombreCampo in arrNombre:
                try:
                    datos[nombreCampo]
                except:
                    valCampo = ""
                else:    
                    valCampo = datos[nombreCampo]
                nuevaRowCell[iC].text = valCampo
                iC=iC+1
    #   Verifica si existe la celda i qeu inicia en 2 si existe remplaza los datos.  La primera borra los campos de combiancion.
        else:
            for nombreCampo in arrNombre:
                try:
                    datos[nombreCampo]
                except:
                    valCampo = ""
                else:    
                    valCampo = datos[nombreCampo]
                tableCombina.rows[i-1].cells[iC].text = valCampo
                iC=iC+1
    i=i+1
# Elimina el campo con el texto tabla_terceros que indica la tabla a usar.
tableCombina.rows[0].cells[0].text = ""
document.save(archivoFinal)
#zairavz96@gmail.com
